from docx import Document
from docx.shared import Cm
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
import os, re, subprocess




def set():
	abspath = os.path.abspath(__file__)
	dname = os.path.dirname(abspath)
	os.chdir(dname)
	# os.chdir("..")
	path = os.path.abspath(os.curdir)

	return path


def Figure(paragraph):
	run = run = paragraph.add_run()
	r = run._r
	fldChar = OxmlElement('w:fldChar')
	fldChar.set(qn('w:fldCharType'), 'begin')
	r.append(fldChar)
	instrText = OxmlElement('w:instrText')
	instrText.text = ' SEQ Figure \* ARABIC'
	r.append(instrText)
	fldChar = OxmlElement('w:fldChar')
	fldChar.set(qn('w:fldCharType'), 'end')
	r.append(fldChar)


def Picture(paragraph):
	run = run = paragraph.add_run()
	r = run._r
	fldChar = OxmlElement('w:fldChar')
	fldChar.set(qn('w:fldCharType'), 'begin')
	r.append(fldChar)
	instrText = OxmlElement('w:instrText')
	instrText.text = ' SEQ Рисунок \* ARABIC'
	r.append(instrText)
	fldChar = OxmlElement('w:fldChar')
	fldChar.set(qn('w:fldCharType'), 'end')
	r.append(fldChar)



def Table(paragraph):
	run = run = paragraph.add_run()
	r = run._r
	fldChar = OxmlElement('w:fldChar')
	fldChar.set(qn('w:fldCharType'), 'begin')
	r.append(fldChar)
	instrText = OxmlElement('w:instrText')
	instrText.text = ' SEQ Table \* ARABIC'
	r.append(instrText)
	fldChar = OxmlElement('w:fldChar')
	fldChar.set(qn('w:fldCharType'), 'end')
	r.append(fldChar)


def Add_Picture(filename, my_caption_str='', my_height=10, my_width=17, my_indent=10):
	document.add_picture(filename, height=Cm(my_height), width=Cm(my_width))            
	paragraph = document.paragraphs[-1]
	paragraph_format = paragraph.paragraph_format
	paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
	paragraph_format.left_indent = Cm(my_indent) # Indenting from 'paragraph.alignment'

	paragraph = document.add_paragraph('Picture ', style='Caption')
	Picture(paragraph)
	paragraph.add_run(my_caption_str)            
	paragraph_format = paragraph.paragraph_format
	paragraph_format.left_indent = Cm(0)
	paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
	paragraph.style = styles # Applying style to paragraph


#
# The main program
#
document = Document()

# Setting style for the document
styles = document.styles['Normal']
styles.paragraph_format.space_before = 0 # No intervals style
styles.paragraph_format.space_after = 0
font = styles.font
font.name = 'Times New Roman'
font.size = Pt(14)


# Sizes of picture in cm
my_height = 5
my_width = 5
# Indent of picture in cm
my_indent = -5



root_dir = set()

working_dir = root_dir + '/_pictures/'
os.chdir(working_dir)

for filename in os.listdir(working_dir):
	if filename.endswith('.png'):
		my_str = ' — a ' + filename.split('_')[0]
		Add_Picture(filename, my_str, my_height, my_width, my_indent)

os.chdir(root_dir)
document.add_page_break()

document.save(str.strip('Figures.docx'))